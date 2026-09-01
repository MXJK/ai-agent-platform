from __future__ import annotations

import ast
from collections import Counter
from dataclasses import replace
from datetime import datetime, timezone
import hashlib
from io import BytesIO
import math
import re
from threading import Lock
from time import perf_counter
from typing import Callable
from uuid import NAMESPACE_URL, uuid4, uuid5

import httpx
from docx import Document
from pypdf import PdfReader

from ai_agent_platform.core import Settings
from ai_agent_platform.integrations.llm import LLMClient
from ai_agent_platform.integrations.rag.errors import (
    RAGConfigurationError,
    RAGError,
    RAGProviderError,
    RAGRerankerUnavailableError,
    RAGValidationError,
)
from ai_agent_platform.integrations.rag.models import (
    DocumentChunk,
    DocumentVectorSnapshot,
    DocumentStore,
    EmbeddingProvider,
    IndexJob,
    IndexJobStore,
    IngestedDocument,
    KnowledgeDocument,
    ParsedDocument,
    RAGAnswer,
    RAGSearchResult,
    Reranker,
    RerankerCapabilities,
    RetrievedDocument,
    RetrievalExecution,
    VectorStore,
)
from ai_agent_platform.usage_ledger import (
    UsageContext,
    current_model_usage_context,
    model_usage_scope,
)


SUPPORTED_TEXT_EXTENSIONS = {
    ".css",
    ".go",
    ".html",
    ".java",
    ".js",
    ".json",
    ".jsx",
    ".md",
    ".markdown",
    ".py",
    ".rs",
    ".toml",
    ".ts",
    ".tsx",
    ".txt",
    ".yaml",
    ".yml",
}
SUPPORTED_DOCUMENT_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | {".docx", ".pdf"}
CODE_TEXT_EXTENSIONS = {
    ".go",
    ".java",
    ".js",
    ".jsx",
    ".py",
    ".rs",
    ".ts",
    ".tsx",
}
INDEX_JOB_TRANSITIONS = {
    "pending": {"parsing", "failed"},
    "parsing": {"embedding", "failed"},
    "embedding": {"vector_written", "failed"},
    "vector_written": {"active", "failed"},
    "active": set(),
    "failed": set(),
}


class InMemoryRAGMetadataStore:
    """Process-local lexical index and index-job journal."""

    def __init__(self) -> None:
        self._documents: dict[str, KnowledgeDocument] = {}
        self._chunks: dict[str, DocumentChunk] = {}
        self._jobs: dict[str, IndexJob] = {}
        self._lock = Lock()

    def save_document(
        self,
        document: ParsedDocument,
        chunks: list[DocumentChunk],
    ) -> KnowledgeDocument:
        with self._lock:
            now = _utcnow()
            existing = self._documents.get(document.id)
            stored = KnowledgeDocument(
                id=document.id,
                knowledge_base_id=document.knowledge_base_id,
                title=(document.title or document.filename).strip(),
                filename=document.filename,
                description=document.description.strip(),
                tags=list(document.tags),
                media_type=document.media_type,
                byte_size=document.byte_size,
                content_hash=hashlib.sha256(document.text.encode("utf-8")).hexdigest(),
                chunk_count=len(chunks),
                is_searchable=True,
                last_index_status="vector_written",
                last_index_error=None,
                created_at=existing.created_at if existing is not None else now,
                updated_at=now,
                indexed_at=now,
                source_uri=document.source_uri,
            )
            self._documents[document.id] = stored
            self._chunks = {
                chunk_id: chunk
                for chunk_id, chunk in self._chunks.items()
                if chunk.document_id != document.id
            }
            self._chunks.update({chunk.id: chunk for chunk in chunks})
            return stored

    def find_document_by_filename(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
    ) -> KnowledgeDocument | None:
        with self._lock:
            for document in self._documents.values():
                if (
                    document.knowledge_base_id == knowledge_base_id
                    and document.filename.casefold() == filename.casefold()
                ):
                    return self._with_latest_job_locked(document)
        return None

    def get_document(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
    ) -> KnowledgeDocument | None:
        with self._lock:
            document = self._documents.get(document_id)
            if document is None or document.knowledge_base_id != knowledge_base_id:
                return None
            return self._with_latest_job_locked(document)

    def list_documents(
        self,
        *,
        knowledge_base_id: str,
        query: str,
        status: str | None,
        sort: str,
        page: int,
        page_size: int,
    ) -> tuple[list[KnowledgeDocument], int]:
        with self._lock:
            documents = [
                self._with_latest_job_locked(document)
                for document in self._documents.values()
                if document.knowledge_base_id == knowledge_base_id
            ]
        query_key = query.strip().casefold()
        if query_key:
            documents = [
                document
                for document in documents
                if query_key
                in " ".join(
                    (
                        document.title,
                        document.filename,
                        document.description,
                        " ".join(document.tags),
                    )
                ).casefold()
            ]
        if status:
            documents = [
                document
                for document in documents
                if document.last_index_status == status
            ]
        if sort == "title_asc":
            documents.sort(key=lambda item: (item.title.casefold(), item.id))
        elif sort == "created_at_desc":
            documents.sort(key=lambda item: (item.created_at, item.id), reverse=True)
        else:
            documents.sort(key=lambda item: (item.updated_at, item.id), reverse=True)
        total = len(documents)
        start = (page - 1) * page_size
        return documents[start : start + page_size], total

    def update_document_metadata(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
        title: str,
        description: str,
        tags: list[str],
    ) -> KnowledgeDocument | None:
        with self._lock:
            document = self._documents.get(document_id)
            if document is None or document.knowledge_base_id != knowledge_base_id:
                return None
            updated = replace(
                document,
                title=title,
                description=description,
                tags=list(tags),
                updated_at=_utcnow(),
            )
            self._documents[document_id] = updated
            return self._with_latest_job_locked(updated)

    def get_document_chunks(
        self,
        *,
        document_id: str,
    ) -> list[DocumentChunk]:
        with self._lock:
            chunks = [
                chunk
                for chunk in self._chunks.values()
                if chunk.document_id == document_id
            ]
        return sorted(chunks, key=lambda chunk: (chunk.chunk_index, chunk.id))

    def delete_document(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
    ) -> bool:
        with self._lock:
            document = self._documents.get(document_id)
            if document is None or document.knowledge_base_id != knowledge_base_id:
                return False
            del self._documents[document_id]
            self._chunks = {
                chunk_id: chunk
                for chunk_id, chunk in self._chunks.items()
                if chunk.document_id != document_id
            }
            return True

    def _with_latest_job_locked(
        self,
        document: KnowledgeDocument,
    ) -> KnowledgeDocument:
        jobs = [
            job
            for job in self._jobs.values()
            if job.document_id == document.id
        ]
        if not jobs:
            return document
        latest = max(jobs, key=lambda job: (job.created_at, job.id))
        return replace(
            document,
            last_index_status=latest.status,
            last_index_error=latest.error,
        )

    def search_lexical(
        self,
        *,
        knowledge_base_id: str,
        query: str,
        limit: int,
    ) -> list[RetrievedDocument]:
        with self._lock:
            chunks = [
                chunk
                for chunk in self._chunks.values()
                if chunk.knowledge_base_id == knowledge_base_id
            ]
        return _bm25_rank_chunks(query=query, chunks=chunks, limit=limit)

    def delete_knowledge_base(self, *, knowledge_base_id: str) -> None:
        with self._lock:
            self._documents = {
                document_id: document
                for document_id, document in self._documents.items()
                if document.knowledge_base_id != knowledge_base_id
            }
            self._chunks = {
                chunk_id: chunk
                for chunk_id, chunk in self._chunks.items()
                if chunk.knowledge_base_id != knowledge_base_id
            }
            self._jobs = {
                job_id: job
                for job_id, job in self._jobs.items()
                if job.knowledge_base_id != knowledge_base_id
            }

    def create_index_job(self, job: IndexJob) -> None:
        with self._lock:
            if job.id in self._jobs:
                raise RAGProviderError(f"index job already exists: {job.id}")
            self._jobs[job.id] = job

    def transition_index_job(
        self,
        *,
        job_id: str,
        expected_status: str,
        status: str,
        document_id: str | None = None,
        chunk_count: int | None = None,
        error: str | None = None,
    ) -> IndexJob:
        with self._lock:
            current = self._jobs.get(job_id)
            if current is None:
                raise RAGProviderError(f"index job not found: {job_id}")
            if current.status != expected_status:
                raise RAGProviderError(
                    f"index job {job_id} expected {expected_status}, "
                    f"found {current.status}"
                )
            if status not in INDEX_JOB_TRANSITIONS[current.status]:
                raise RAGProviderError(
                    f"invalid index transition: {current.status} -> {status}"
                )
            now = _utcnow()
            updated = replace(
                current,
                status=status,
                document_id=document_id or current.document_id,
                chunk_count=(
                    chunk_count if chunk_count is not None else current.chunk_count
                ),
                error=error,
                updated_at=now,
                completed_at=now if status in {"active", "failed"} else None,
            )
            self._jobs[job_id] = updated
            return updated

    def get_index_job(self, job_id: str) -> IndexJob | None:
        with self._lock:
            return self._jobs.get(job_id)

    def list_index_jobs(
        self,
        *,
        knowledge_base_id: str,
        limit: int,
    ) -> list[IndexJob]:
        with self._lock:
            jobs = [
                job
                for job in self._jobs.values()
                if job.knowledge_base_id == knowledge_base_id
            ]
        jobs.sort(key=lambda job: (job.created_at, job.id), reverse=True)
        return jobs[:limit]


class TextDocumentParser:
    """Extracts and normalizes supported text, PDF, and DOCX documents."""

    def parse(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
        content: str,
        source_uri: str | None = None,
        document_id: str | None = None,
        title: str | None = None,
        description: str = "",
        tags: list[str] | None = None,
        media_type: str | None = None,
        byte_size: int | None = None,
    ) -> ParsedDocument:
        extension = _file_extension(filename)
        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
            raise RAGValidationError(
                f"unsupported document type: {extension or '<none>'}; "
                f"supported types: {supported}"
            )

        if extension in CODE_TEXT_EXTENSIONS:
            text = _normalize_code_text(content)
        else:
            text = _normalize_text(content)
        if not text:
            raise RAGValidationError("document text is empty")

        return ParsedDocument(
            id=document_id
            or _document_id(
                    knowledge_base_id=knowledge_base_id,
                    filename=filename,
                    source_uri=source_uri,
                ),
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            text=text,
            source_uri=source_uri,
            title=title,
            description=description,
            tags=list(tags or []),
            media_type=media_type,
            byte_size=byte_size,
        )

    def parse_bytes(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
        content: bytes,
        source_uri: str | None = None,
        document_id: str | None = None,
        title: str | None = None,
        description: str = "",
        tags: list[str] | None = None,
        media_type: str | None = None,
        byte_size: int | None = None,
    ) -> ParsedDocument:
        extension = _file_extension(filename)
        if extension in SUPPORTED_TEXT_EXTENSIONS:
            try:
                text = content.decode("utf-8-sig")
            except UnicodeDecodeError as exc:
                raise RAGValidationError(
                    "text documents must use UTF-8 encoding"
                ) from exc
        elif extension == ".pdf":
            text = _extract_pdf_text(content)
        elif extension == ".docx":
            text = _extract_docx_text(content)
        else:
            supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
            raise RAGValidationError(
                f"unsupported document type: {extension or '<none>'}; "
                f"supported types: {supported}"
            )
        return self.parse(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            content=text,
            source_uri=source_uri,
            document_id=document_id,
            title=title,
            description=description,
            tags=tags,
            media_type=media_type,
            byte_size=byte_size if byte_size is not None else len(content),
        )


class RecursiveCharacterChunker:
    def __init__(self, *, chunk_size: int = 800, chunk_overlap: int = 120) -> None:
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")
        if chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be smaller than chunk_size")
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap

    def split(self, document: ParsedDocument) -> list[DocumentChunk]:
        if _file_extension(document.filename) in CODE_TEXT_EXTENSIONS:
            code_chunks = self._split_code_document(document)
            if code_chunks:
                return code_chunks
            return self._split_text_with_line_ranges(document)
        return self._split_text(document)

    def _split_text(self, document: ParsedDocument) -> list[DocumentChunk]:
        text = document.text
        chunks: list[DocumentChunk] = []
        start = 0
        text_length = len(text)

        while start < text_length:
            raw_end = min(start + self._chunk_size, text_length)
            end = self._best_breakpoint(text, start, raw_end)
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    DocumentChunk(
                        id=_chunk_id(
                            document_id=document.id,
                            chunk_index=len(chunks),
                        ),
                        knowledge_base_id=document.knowledge_base_id,
                        document_id=document.id,
                        filename=document.filename,
                        chunk_index=len(chunks),
                        text=chunk_text,
                        start_line=_line_number_for_offset(text, start),
                        end_line=_line_number_for_offset(text, end),
                    )
                )
            if end >= text_length:
                break
            start = max(end - self._chunk_overlap, start + 1)

        return chunks

    def _split_code_document(self, document: ParsedDocument) -> list[DocumentChunk]:
        lines = document.text.splitlines()
        starts = _code_symbol_starts(lines, document.filename)
        chunks: list[DocumentChunk] = []

        if starts and starts[0][0] > 1:
            preamble_text = "\n".join(lines[: starts[0][0] - 1]).strip()
            if preamble_text:
                chunks.extend(
                    self._chunks_from_line_block(
                        document=document,
                        lines=lines[: starts[0][0] - 1],
                        start_line=1,
                        symbols=[],
                        start_index=len(chunks),
                    )
                )

        for index, (start_line, symbols) in enumerate(starts):
            next_start = starts[index + 1][0] if index + 1 < len(starts) else len(lines) + 1
            block_lines = lines[start_line - 1 : next_start - 1]
            chunks.extend(
                self._chunks_from_line_block(
                    document=document,
                    lines=block_lines,
                    start_line=start_line,
                    symbols=symbols,
                    start_index=len(chunks),
                )
            )
        return chunks

    def _split_text_with_line_ranges(self, document: ParsedDocument) -> list[DocumentChunk]:
        lines = document.text.splitlines()
        return self._chunks_from_line_block(
            document=document,
            lines=lines,
            start_line=1,
            symbols=[],
            start_index=0,
        )

    def _chunks_from_line_block(
        self,
        *,
        document: ParsedDocument,
        lines: list[str],
        start_line: int,
        symbols: list[str],
        start_index: int,
    ) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []
        current_lines: list[str] = []
        current_start_line = start_line
        current_chars = 0

        for offset, line in enumerate(lines):
            line_length = len(line) + 1
            if current_lines and current_chars + line_length > self._chunk_size:
                chunks.append(
                    self._build_line_chunk(
                        document=document,
                        chunk_index=start_index + len(chunks),
                        lines=current_lines,
                        start_line=current_start_line,
                        symbols=symbols,
                    )
                )
                overlap_lines = _line_overlap(current_lines, self._chunk_overlap)
                current_start_line = start_line + offset - len(overlap_lines)
                current_lines = overlap_lines
                current_chars = sum(len(item) + 1 for item in current_lines)
            current_lines.append(line)
            current_chars += line_length

        if current_lines and "\n".join(current_lines).strip():
            chunks.append(
                self._build_line_chunk(
                    document=document,
                    chunk_index=start_index + len(chunks),
                    lines=current_lines,
                    start_line=current_start_line,
                    symbols=symbols,
                )
            )
        return chunks

    def _build_line_chunk(
        self,
        *,
        document: ParsedDocument,
        chunk_index: int,
        lines: list[str],
        start_line: int,
        symbols: list[str],
    ) -> DocumentChunk:
        chunk_text = "\n".join(lines).strip()
        return DocumentChunk(
            id=_chunk_id(
                document_id=document.id,
                chunk_index=chunk_index,
            ),
            knowledge_base_id=document.knowledge_base_id,
            document_id=document.id,
            filename=document.filename,
            chunk_index=chunk_index,
            text=chunk_text,
            start_line=start_line,
            end_line=start_line + len(lines) - 1,
            symbols=symbols,
        )

    def _best_breakpoint(self, text: str, start: int, raw_end: int) -> int:
        if raw_end >= len(text):
            return raw_end

        min_end = start + max(int(self._chunk_size * 0.6), 1)
        candidates = [
            text.rfind("\n\n", start, raw_end),
            text.rfind("\n", start, raw_end),
            text.rfind("。", start, raw_end),
            text.rfind(".", start, raw_end),
            text.rfind(" ", start, raw_end),
        ]
        for candidate in candidates:
            if candidate >= min_end:
                return candidate + 1
        return raw_end


class HashingEmbeddingProvider:
    """Small deterministic embedding provider for local learning and tests.

    It is not a semantic model, but it gives stable vectors and lets the RAG
    pipeline run without external API keys. Swap to OpenAIEmbeddingProvider for
    real semantic retrieval quality.
    """

    def __init__(
        self,
        *,
        dimensions: int = 128,
        model: str = "local-hashing",
        usage_ledger=None,
    ) -> None:
        self._dimensions = dimensions
        self._model = model
        self._usage_ledger = usage_ledger

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        embeddings = [self._embed(text) for text in texts]
        if self._usage_ledger is not None:
            self._usage_ledger.record(
                provider="local",
                model=self._model,
                input_tokens=sum(_embedding_text_tokens(text) for text in texts),
                output_tokens=0,
                input_count_method="local_lexical_tokenizer",
                context=_embedding_usage_context(),
            )
        return embeddings

    def _embed(self, text: str) -> list[float]:
        vector = [0.0] * self._dimensions
        tokens = _tokenize(text)
        if not tokens:
            return vector

        for token in tokens:
            digest = hashlib.sha256(token.encode("utf-8")).digest()
            index = int.from_bytes(digest[:4], "big") % self._dimensions
            sign = 1.0 if digest[4] % 2 == 0 else -1.0
            vector[index] += sign

        norm = math.sqrt(sum(value * value for value in vector))
        if norm == 0:
            return vector
        return [value / norm for value in vector]


class SentenceTransformerEmbeddingProvider:
    """Lazy local semantic embeddings backed by sentence-transformers."""

    def __init__(
        self,
        *,
        model_name: str,
        device: str = "cpu",
        usage_ledger=None,
    ) -> None:
        model_name = model_name.strip()
        device = device.strip()
        if not model_name:
            raise ValueError("model_name must not be empty")
        if not device:
            raise ValueError("device must not be empty")
        self.model_name = model_name
        self.device = device
        self._usage_ledger = usage_ledger
        self._model = None
        self._model_lock = Lock()
        self._load_error: Exception | None = None

    @property
    def status(self) -> str:
        if self._model is not None:
            return "ready"
        if self._load_error is not None:
            return "error"
        return "not_loaded"

    def _get_model(self):
        if self._model is not None:
            return self._model
        with self._model_lock:
            if self._model is not None:
                return self._model
            if self._load_error is not None:
                raise RAGProviderError(
                    "sentence-transformer embedding model failed to load"
                ) from self._load_error
            try:
                from sentence_transformers import SentenceTransformer

                self._model = SentenceTransformer(
                    self.model_name,
                    device=self.device,
                )
            except ImportError as exc:
                self._load_error = exc
                raise RAGConfigurationError(
                    "sentence-transformers is not installed; install project dependencies"
                ) from exc
            except Exception as exc:
                self._load_error = exc
                raise RAGProviderError(
                    "sentence-transformer embedding model failed to load"
                ) from exc
        return self._model

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        if not texts:
            return []
        try:
            encoded = self._get_model().encode(
                texts,
                batch_size=32,
                show_progress_bar=False,
                convert_to_numpy=True,
                normalize_embeddings=True,
            )
            embeddings = [
                [float(value) for value in vector]
                for vector in encoded
            ]
        except (RAGConfigurationError, RAGProviderError):
            raise
        except Exception as exc:
            raise RAGProviderError(
                "sentence-transformer embedding inference failed"
            ) from exc
        if len(embeddings) != len(texts):
            raise RAGProviderError(
                "sentence-transformer embedding model returned malformed data"
            )
        if self._usage_ledger is not None:
            self._usage_ledger.record(
                provider="sentence_transformer",
                model=self.model_name,
                input_tokens=sum(_embedding_text_tokens(text) for text in texts),
                output_tokens=0,
                input_count_method="local_lexical_tokenizer_estimate",
                context=_embedding_usage_context(),
            )
        return embeddings


class OpenAIEmbeddingProvider:
    def __init__(
        self,
        settings: Settings,
        usage_ledger=None,
        credential_resolver: Callable[[str], str | None] | None = None,
    ) -> None:
        self._credential_resolver = credential_resolver
        self._model = settings.embedding_model
        self._timeout_seconds = settings.llm_timeout_seconds
        self._usage_ledger = usage_ledger

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        api_key = (
            self._credential_resolver("openai")
            if self._credential_resolver is not None
            else None
        )
        if not api_key:
            raise RAGConfigurationError(
                "OpenAI credential is required in model management for embeddings"
            )
        payload = {
            "model": self._model,
            "input": texts,
        }
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
        timeout = httpx.Timeout(self._timeout_seconds)
        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                "https://api.openai.com/v1/embeddings",
                headers=headers,
                json=payload,
            )
        if response.status_code >= 400:
            raise RAGProviderError(
                f"embedding provider returned HTTP {response.status_code}"
            )

        body = response.json()
        data = body.get("data")
        if not isinstance(data, list):
            raise RAGProviderError("embedding provider returned malformed data")
        usage = body.get("usage")
        if not isinstance(usage, dict) or not isinstance(
            usage.get("prompt_tokens"),
            int,
        ):
            raise RAGProviderError("embedding provider returned no token usage")
        if self._usage_ledger is not None:
            self._usage_ledger.record(
                provider="openai",
                model=self._model,
                input_tokens=max(0, int(usage["prompt_tokens"])),
                output_tokens=0,
                input_count_method="openai_embedding_usage",
                context=_embedding_usage_context(),
            )
        return [item["embedding"] for item in data]


class GeminiEmbeddingProvider:
    def __init__(
        self,
        settings: Settings,
        usage_ledger=None,
        credential_resolver: Callable[[str], str | None] | None = None,
    ) -> None:
        self._credential_resolver = credential_resolver
        self._model = settings.embedding_model
        self._timeout_seconds = settings.llm_timeout_seconds
        self._usage_ledger = usage_ledger

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        api_key = (
            self._credential_resolver("google")
            if self._credential_resolver is not None
            else None
        )
        if not api_key:
            raise RAGConfigurationError(
                "Google credential is required in model management for Gemini embeddings"
            )

        gemini_task_type = _gemini_task_type(task_type)
        return [
            self._embed_one(
                text=text,
                task_type=gemini_task_type,
                api_key=api_key,
            )
            for text in texts
        ]

    def _embed_one(
        self,
        *,
        text: str,
        task_type: str,
        api_key: str,
    ) -> list[float]:
        payload = {
            "content": {
                "parts": [
                    {
                        "text": text,
                    }
                ],
            },
            "taskType": task_type,
        }
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": api_key,
        }
        model = self._model
        if not model.startswith("models/"):
            model = f"models/{model}"
        timeout = httpx.Timeout(self._timeout_seconds)
        with httpx.Client(timeout=timeout) as client:
            count_response = client.post(
                (
                    "https://generativelanguage.googleapis.com/v1beta/"
                    f"{model}:countTokens"
                ),
                headers=headers,
                json={"contents": [payload["content"]]},
            )
            if count_response.status_code >= 400:
                raise RAGProviderError(
                    "Gemini token count provider returned HTTP "
                    f"{count_response.status_code}"
                )
            count_body = count_response.json()
            input_tokens = count_body.get("totalTokens")
            if not isinstance(input_tokens, int):
                raise RAGProviderError(
                    "Gemini token count provider returned malformed data"
                )
            response = client.post(
                f"https://generativelanguage.googleapis.com/v1beta/{model}:embedContent",
                headers=headers,
                json=payload,
            )
        if response.status_code >= 400:
            raise RAGProviderError(
                f"Gemini embedding provider returned HTTP {response.status_code}"
            )

        body = response.json()
        embedding = body.get("embedding")
        if not isinstance(embedding, dict):
            raise RAGProviderError("Gemini embedding provider returned malformed data")
        values = embedding.get("values")
        if not isinstance(values, list):
            raise RAGProviderError("Gemini embedding provider returned no values")
        if self._usage_ledger is not None:
            self._usage_ledger.record(
                provider="gemini",
                model=self._model,
                input_tokens=max(0, input_tokens),
                output_tokens=0,
                input_count_method="gemini_models_count_tokens",
                context=_embedding_usage_context(),
            )
        return [float(value) for value in values]


class NoopReranker:
    status = "unavailable"

    def rerank(
        self,
        *,
        query: str,
        candidates: list[RetrievedDocument],
        limit: int,
    ) -> list[RetrievedDocument]:
        return candidates[:limit]


class SentenceTransformerCrossEncoderReranker:
    def __init__(self, *, model_name: str, device: str = "cpu") -> None:
        device = device.strip()
        if not device:
            raise ValueError("device must not be empty")
        self.model_name = model_name
        self.device = device
        self._model = None
        self._model_lock = Lock()
        self._load_error: Exception | None = None

    @property
    def status(self) -> str:
        if self._model is not None:
            return "ready"
        if self._load_error is not None:
            return "error"
        return "not_loaded"

    def _get_model(self):
        if self._model is not None:
            return self._model
        with self._model_lock:
            if self._model is not None:
                return self._model
            if self._load_error is not None:
                raise RAGProviderError(
                    "reranker model failed to load"
                ) from self._load_error
            try:
                from sentence_transformers import CrossEncoder

                self._model = CrossEncoder(self.model_name, device=self.device)
            except ImportError as exc:
                self._load_error = exc
                raise RAGConfigurationError(
                    "sentence-transformers is not installed; "
                    "run pip install -r requirements.txt"
                ) from exc
            except Exception as exc:
                self._load_error = exc
                raise RAGProviderError("reranker model failed to load") from exc
        return self._model

    def rerank(
        self,
        *,
        query: str,
        candidates: list[RetrievedDocument],
        limit: int,
    ) -> list[RetrievedDocument]:
        if not candidates:
            return []

        pairs = [(query, candidate.text) for candidate in candidates]
        try:
            raw_scores = self._get_model().predict(pairs)
        except (RAGConfigurationError, RAGProviderError):
            raise
        except Exception as exc:
            raise RAGProviderError("reranker inference failed") from exc
        scored = [
            replace(
                candidate,
                score=float(score),
                recall_score=(
                    candidate.recall_score
                    if candidate.recall_score is not None
                    else candidate.score
                ),
                rerank_score=float(score),
            )
            for candidate, score in zip(candidates, raw_scores)
        ]
        scored.sort(key=lambda item: item.score, reverse=True)
        return scored[:limit]


class InMemoryVectorStore:
    def __init__(self) -> None:
        self._rows: list[tuple[DocumentChunk, list[float]]] = []
        self._lock = Lock()

    def delete_document(self, *, document_id: str) -> None:
        with self._lock:
            self._rows = [
                (chunk, embedding)
                for chunk, embedding in self._rows
                if chunk.document_id != document_id
            ]

    def delete_knowledge_base(self, *, knowledge_base_id: str) -> None:
        with self._lock:
            self._rows = [
                (chunk, embedding)
                for chunk, embedding in self._rows
                if chunk.knowledge_base_id != knowledge_base_id
            ]

    def snapshot_document(
        self,
        *,
        document_id: str,
    ) -> DocumentVectorSnapshot:
        with self._lock:
            rows = [
                (chunk, list(embedding))
                for chunk, embedding in self._rows
                if chunk.document_id == document_id
            ]
        return DocumentVectorSnapshot(
            chunks=[chunk for chunk, _ in rows],
            embeddings=[embedding for _, embedding in rows],
        )

    def restore_document(
        self,
        *,
        document_id: str,
        snapshot: DocumentVectorSnapshot,
    ) -> None:
        self.replace_document(
            document_id=document_id,
            chunks=snapshot.chunks,
            embeddings=snapshot.embeddings,
        )

    def upsert_chunks(
        self,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        _validate_embeddings(chunks, embeddings)
        with self._lock:
            existing_ids = {chunk.id for chunk in chunks}
            self._rows = [
                (chunk, embedding)
                for chunk, embedding in self._rows
                if chunk.id not in existing_ids
            ]
            self._rows.extend(zip(chunks, embeddings))

    def replace_document(
        self,
        *,
        document_id: str,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        _validate_embeddings(chunks, embeddings)
        with self._lock:
            retained = [
                (chunk, embedding)
                for chunk, embedding in self._rows
                if chunk.document_id != document_id
            ]
            self._rows = retained + list(zip(chunks, embeddings))

    def search(
        self,
        *,
        knowledge_base_id: str,
        query_embedding: list[float],
        limit: int,
    ) -> list[RetrievedDocument]:
        scored: list[RetrievedDocument] = []
        with self._lock:
            rows = list(self._rows)
        for chunk, embedding in rows:
            if chunk.knowledge_base_id != knowledge_base_id:
                continue
            score = _cosine_similarity(query_embedding, embedding)
            scored.append(
                RetrievedDocument(
                    id=chunk.id,
                    knowledge_base_id=chunk.knowledge_base_id,
                    document_id=chunk.document_id,
                    filename=chunk.filename,
                    chunk_index=chunk.chunk_index,
                    text=chunk.text,
                    score=score,
                    start_line=chunk.start_line,
                    end_line=chunk.end_line,
                    symbols=chunk.symbols,
                    recall_score=score,
                )
            )
        scored.sort(key=lambda item: item.score, reverse=True)
        return scored[:limit]


class ChromaVectorStore:
    def __init__(self, *, persist_directory: str, collection_name: str) -> None:
        try:
            import chromadb
        except ImportError as exc:
            raise RAGConfigurationError(
                "chromadb is not installed; run pip install -r requirements.txt"
            ) from exc

        client = chromadb.PersistentClient(path=persist_directory)
        self._collection = client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

    def delete_document(self, *, document_id: str) -> None:
        self._collection.delete(where={"document_id": document_id})

    def delete_knowledge_base(self, *, knowledge_base_id: str) -> None:
        self._collection.delete(where={"knowledge_base_id": knowledge_base_id})

    def snapshot_document(
        self,
        *,
        document_id: str,
    ) -> DocumentVectorSnapshot:
        result = self._collection.get(
            where={"document_id": document_id},
            include=["documents", "metadatas", "embeddings"],
        )
        ids = list(result.get("ids") or [])
        documents = list(result.get("documents") or [])
        metadatas = list(result.get("metadatas") or [])
        raw_embeddings = result.get("embeddings")
        embeddings = list(raw_embeddings) if raw_embeddings is not None else []
        chunks = [
            DocumentChunk(
                id=str(chunk_id),
                knowledge_base_id=str(metadata["knowledge_base_id"]),
                document_id=str(metadata["document_id"]),
                filename=str(metadata["filename"]),
                chunk_index=int(metadata["chunk_index"]),
                text=str(documents[index]),
                start_line=_optional_int(metadata.get("start_line")),
                end_line=_optional_int(metadata.get("end_line")),
                symbols=_metadata_symbols(metadata.get("symbols")),
            )
            for index, (chunk_id, metadata) in enumerate(zip(ids, metadatas))
        ]
        return DocumentVectorSnapshot(
            chunks=chunks,
            embeddings=[[float(value) for value in row] for row in embeddings],
        )

    def restore_document(
        self,
        *,
        document_id: str,
        snapshot: DocumentVectorSnapshot,
    ) -> None:
        self.replace_document(
            document_id=document_id,
            chunks=snapshot.chunks,
            embeddings=snapshot.embeddings,
        )

    def upsert_chunks(
        self,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        if not chunks:
            return
        _validate_embeddings(chunks, embeddings)

        self._collection.upsert(
            ids=[chunk.id for chunk in chunks],
            embeddings=embeddings,
            documents=[chunk.text for chunk in chunks],
            metadatas=[_chroma_chunk_metadata(chunk) for chunk in chunks],
        )

    def replace_document(
        self,
        *,
        document_id: str,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        _validate_embeddings(chunks, embeddings)
        existing = self._collection.get(
            where={"document_id": document_id},
            include=[],
        )
        existing_ids = {
            str(item) for item in (existing.get("ids") or [])
        }
        self.upsert_chunks(chunks, embeddings)
        stale_ids = existing_ids - {chunk.id for chunk in chunks}
        if stale_ids:
            self._collection.delete(ids=sorted(stale_ids))

    def search(
        self,
        *,
        knowledge_base_id: str,
        query_embedding: list[float],
        limit: int,
    ) -> list[RetrievedDocument]:
        result = self._collection.query(
            query_embeddings=[query_embedding],
            n_results=limit,
            where={"knowledge_base_id": knowledge_base_id},
            include=["documents", "metadatas", "distances"],
        )

        ids = _first_result_list(result.get("ids"))
        documents = _first_result_list(result.get("documents"))
        metadatas = _first_result_list(result.get("metadatas"))
        distances = _first_result_list(result.get("distances"))

        retrieved: list[RetrievedDocument] = []
        for index, chunk_id in enumerate(ids):
            metadata = metadatas[index]
            distance = distances[index]
            score = 1.0 - float(distance)
            retrieved.append(
                RetrievedDocument(
                    id=str(chunk_id),
                    knowledge_base_id=str(metadata["knowledge_base_id"]),
                    document_id=str(metadata["document_id"]),
                    filename=str(metadata["filename"]),
                    chunk_index=int(metadata["chunk_index"]),
                    text=str(documents[index]),
                    score=score,
                    start_line=_optional_int(metadata.get("start_line")),
                    end_line=_optional_int(metadata.get("end_line")),
                    symbols=_metadata_symbols(metadata.get("symbols")),
                    recall_score=score,
                )
            )
        return retrieved


class QdrantVectorStore:
    def __init__(
        self,
        *,
        url: str,
        api_key: str | None,
        collection_name: str,
    ) -> None:
        try:
            from qdrant_client import QdrantClient
        except ImportError as exc:
            raise RAGConfigurationError(
                "qdrant-client is not installed; run pip install -r requirements.txt"
            ) from exc

        if url == ":memory:":
            self._client = QdrantClient(location=":memory:")
        else:
            self._client = QdrantClient(url=url, api_key=api_key)
        self._collection_name = collection_name
        self._vector_size: int | None = None

    def delete_document(self, *, document_id: str) -> None:
        if not self._collection_exists():
            return
        FieldCondition = _qdrant_model("FieldCondition")
        Filter = _qdrant_model("Filter")
        MatchValue = _qdrant_model("MatchValue")
        FilterSelector = _qdrant_model("FilterSelector")
        self._client.delete(
            collection_name=self._collection_name,
            points_selector=FilterSelector(
                filter=Filter(
                    must=[
                        FieldCondition(
                            key="document_id",
                            match=MatchValue(value=document_id),
                        )
                    ]
                )
            ),
        )

    def delete_knowledge_base(self, *, knowledge_base_id: str) -> None:
        if not self._collection_exists():
            return
        FieldCondition = _qdrant_model("FieldCondition")
        Filter = _qdrant_model("Filter")
        MatchValue = _qdrant_model("MatchValue")
        FilterSelector = _qdrant_model("FilterSelector")
        self._client.delete(
            collection_name=self._collection_name,
            points_selector=FilterSelector(
                filter=Filter(
                    must=[
                        FieldCondition(
                            key="knowledge_base_id",
                            match=MatchValue(value=knowledge_base_id),
                        )
                    ]
                )
            ),
        )

    def snapshot_document(
        self,
        *,
        document_id: str,
    ) -> DocumentVectorSnapshot:
        if not self._collection_exists():
            return DocumentVectorSnapshot(chunks=[], embeddings=[])
        points = self._document_points(
            document_id=document_id,
            with_payload=True,
            with_vectors=True,
        )
        chunks: list[DocumentChunk] = []
        embeddings: list[list[float]] = []
        for point in points:
            payload = point.payload or {}
            vector = point.vector
            if isinstance(vector, dict):
                vector = next(iter(vector.values()), [])
            chunks.append(
                DocumentChunk(
                    id=str(payload["chunk_id"]),
                    knowledge_base_id=str(payload["knowledge_base_id"]),
                    document_id=str(payload["document_id"]),
                    filename=str(payload["filename"]),
                    chunk_index=int(payload["chunk_index"]),
                    text=str(payload["text"]),
                    start_line=_optional_int(payload.get("start_line")),
                    end_line=_optional_int(payload.get("end_line")),
                    symbols=_metadata_symbols(payload.get("symbols")),
                )
            )
            embeddings.append([float(value) for value in (vector or [])])
        return DocumentVectorSnapshot(chunks=chunks, embeddings=embeddings)

    def restore_document(
        self,
        *,
        document_id: str,
        snapshot: DocumentVectorSnapshot,
    ) -> None:
        self.replace_document(
            document_id=document_id,
            chunks=snapshot.chunks,
            embeddings=snapshot.embeddings,
        )

    def upsert_chunks(
        self,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        if not chunks:
            return
        _validate_embeddings(chunks, embeddings)

        self._ensure_collection(vector_size=len(embeddings[0]))
        PointStruct = _qdrant_model("PointStruct")
        points = [
            PointStruct(
                id=_qdrant_point_id(chunk.id),
                vector=embedding,
                payload={
                    "chunk_id": chunk.id,
                    "knowledge_base_id": chunk.knowledge_base_id,
                    "document_id": chunk.document_id,
                    "filename": chunk.filename,
                    "chunk_index": chunk.chunk_index,
                    "start_line": chunk.start_line,
                    "end_line": chunk.end_line,
                    "symbols": chunk.symbols,
                    "text": chunk.text,
                },
            )
            for chunk, embedding in zip(chunks, embeddings)
        ]
        self._client.upsert(collection_name=self._collection_name, points=points)

    def replace_document(
        self,
        *,
        document_id: str,
        chunks: list[DocumentChunk],
        embeddings: list[list[float]],
    ) -> None:
        _validate_embeddings(chunks, embeddings)
        existing_ids = self._document_point_ids(document_id=document_id)
        self.upsert_chunks(chunks, embeddings)
        current_ids = {_qdrant_point_id(chunk.id) for chunk in chunks}
        stale_ids = existing_ids - current_ids
        if stale_ids:
            self._client.delete(
                collection_name=self._collection_name,
                points_selector=sorted(stale_ids),
            )

    def search(
        self,
        *,
        knowledge_base_id: str,
        query_embedding: list[float],
        limit: int,
    ) -> list[RetrievedDocument]:
        if not self._collection_exists():
            return []

        FieldCondition = _qdrant_model("FieldCondition")
        Filter = _qdrant_model("Filter")
        MatchValue = _qdrant_model("MatchValue")
        query_filter = Filter(
            must=[
                FieldCondition(
                    key="knowledge_base_id",
                    match=MatchValue(value=knowledge_base_id),
                )
            ]
        )
        if hasattr(self._client, "search"):
            results = self._client.search(
                collection_name=self._collection_name,
                query_vector=query_embedding,
                query_filter=query_filter,
                limit=limit,
                with_payload=True,
            )
        else:
            query_response = self._client.query_points(
                collection_name=self._collection_name,
                query=query_embedding,
                query_filter=query_filter,
                limit=limit,
                with_payload=True,
            )
            results = query_response.points

        retrieved: list[RetrievedDocument] = []
        for point in results:
            payload = point.payload or {}
            retrieved.append(
                RetrievedDocument(
                    id=str(payload["chunk_id"]),
                    knowledge_base_id=str(payload["knowledge_base_id"]),
                    document_id=str(payload["document_id"]),
                    filename=str(payload["filename"]),
                    chunk_index=int(payload["chunk_index"]),
                    text=str(payload["text"]),
                    score=float(point.score),
                    start_line=_optional_int(payload.get("start_line")),
                    end_line=_optional_int(payload.get("end_line")),
                    symbols=_metadata_symbols(payload.get("symbols")),
                    recall_score=float(point.score),
                )
            )
        return retrieved

    def _ensure_collection(self, *, vector_size: int) -> None:
        if self._vector_size is not None:
            if self._vector_size != vector_size:
                raise RAGConfigurationError(
                    "Qdrant collection vector size does not match embedding size"
                )
            return

        if self._collection_exists():
            info = self._client.get_collection(collection_name=self._collection_name)
            configured_size = _qdrant_vector_size(info)
            if configured_size is not None and configured_size != vector_size:
                raise RAGConfigurationError(
                    "Qdrant collection vector size does not match embedding size"
                )
            self._vector_size = configured_size or vector_size
            return

        Distance = _qdrant_model("Distance")
        VectorParams = _qdrant_model("VectorParams")
        self._client.create_collection(
            collection_name=self._collection_name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )
        self._vector_size = vector_size

    def _collection_exists(self) -> bool:
        collections = self._client.get_collections().collections
        return any(item.name == self._collection_name for item in collections)

    def _document_point_ids(self, *, document_id: str) -> set[str]:
        return {
            str(point.id)
            for point in self._document_points(
                document_id=document_id,
                with_payload=False,
                with_vectors=False,
            )
        }

    def _document_points(
        self,
        *,
        document_id: str,
        with_payload: bool,
        with_vectors: bool,
    ) -> list[object]:
        if not self._collection_exists():
            return []
        FieldCondition = _qdrant_model("FieldCondition")
        Filter = _qdrant_model("Filter")
        MatchValue = _qdrant_model("MatchValue")
        document_filter = Filter(
            must=[
                FieldCondition(
                    key="document_id",
                    match=MatchValue(value=document_id),
                )
            ]
        )
        collected: list[object] = []
        offset = None
        while True:
            points, next_offset = self._client.scroll(
                collection_name=self._collection_name,
                scroll_filter=document_filter,
                limit=256,
                offset=offset,
                with_payload=with_payload,
                with_vectors=with_vectors,
            )
            collected.extend(points)
            if next_offset is None:
                return collected
            offset = next_offset


class RAGService:
    def __init__(
        self,
        *,
        parser: TextDocumentParser,
        chunker: RecursiveCharacterChunker,
        embedding_provider: EmbeddingProvider,
        vector_store: VectorStore,
        reranker: Reranker,
        default_recall_limit: int,
        max_prompt_chars: int,
        lexical_weight: float = 0.35,
        document_store: DocumentStore | None = None,
        index_job_store: IndexJobStore | None = None,
        rrf_k: int = 60,
        reranker_provider: str | None = None,
        reranker_model: str | None = None,
        rerank_default_enabled: bool | None = None,
    ) -> None:
        if not 0.0 <= lexical_weight <= 1.0:
            raise ValueError("lexical_weight must be between 0 and 1")
        if rrf_k <= 0:
            raise ValueError("rrf_k must be positive")
        self._parser = parser
        self._chunker = chunker
        self._embedding_provider = embedding_provider
        self._vector_store = vector_store
        self._reranker = reranker
        self._reranker_available = not isinstance(reranker, NoopReranker)
        self._reranker_provider = (
            reranker_provider if self._reranker_available else None
        )
        self._reranker_model = reranker_model if self._reranker_available else None
        self._rerank_default_enabled = (
            self._reranker_available
            if rerank_default_enabled is None
            else rerank_default_enabled
        )
        if self._rerank_default_enabled and not self._reranker_available:
            raise ValueError("rerank_default_enabled requires a configured reranker")
        self._default_recall_limit = default_recall_limit
        self._max_prompt_chars = max_prompt_chars
        self._lexical_weight = lexical_weight
        self._rrf_k = rrf_k
        self._memory_metadata_store = InMemoryRAGMetadataStore()
        self._document_store = document_store or self._memory_metadata_store
        self._index_job_store = (
            index_job_store
            or (
                self._document_store
                if _supports_index_job_store(self._document_store)
                else self._memory_metadata_store
            )
        )

    def ingest_document(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
        content: str,
        source_uri: str | None = None,
        title: str | None = None,
        description: str = "",
        tags: list[str] | None = None,
        media_type: str | None = None,
        byte_size: int | None = None,
    ) -> IngestedDocument:
        return self._run_index_job(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            parse_document=lambda: self._parser.parse(
                knowledge_base_id=knowledge_base_id,
                filename=filename,
                content=content,
                source_uri=source_uri,
                title=title,
                description=description,
                tags=tags,
                media_type=media_type,
                byte_size=byte_size,
            ),
        )

    def _run_index_job(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
        parse_document,
        replacing_document_id: str | None = None,
    ) -> IngestedDocument:
        now = _utcnow()
        job = IndexJob(
            id=f"idx_{uuid4().hex[:16]}",
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            status="pending",
            document_id=replacing_document_id,
            chunk_count=0,
            error=None,
            created_at=now,
            updated_at=now,
        )
        self._index_job_store.create_index_job(job)
        current_status = "pending"
        document: ParsedDocument | None = None
        chunks: list[DocumentChunk] = []
        vector_snapshot: DocumentVectorSnapshot | None = None
        vector_mutated = False
        stored_document: KnowledgeDocument | None = None
        try:
            job = self._transition_index_job(
                job=job,
                status="parsing",
            )
            current_status = job.status
            document = parse_document()
            chunks = self._chunker.split(document)
            if not chunks:
                raise RAGValidationError("document produced no chunks")
            job = self._transition_index_job(
                job=job,
                status="embedding",
                document_id=document.id,
                chunk_count=len(chunks),
            )
            current_status = job.status
            with model_usage_scope(
                operation="embedding",
                resource_id=knowledge_base_id,
            ):
                embeddings = self._embedding_provider.embed_texts(
                    [chunk.text for chunk in chunks],
                    task_type="document",
                )
            _validate_embeddings(chunks, embeddings)
            if replacing_document_id is not None:
                vector_snapshot = self._vector_store.snapshot_document(
                    document_id=replacing_document_id
                )
            replace_document = getattr(
                self._vector_store,
                "replace_document",
                None,
            )
            vector_mutated = True
            if callable(replace_document):
                replace_document(
                    document_id=document.id,
                    chunks=chunks,
                    embeddings=embeddings,
                )
            else:
                self._vector_store.upsert_chunks(chunks, embeddings)
            job = self._transition_index_job(
                job=job,
                status="vector_written",
            )
            current_status = job.status
            stored_document = self._document_store.save_document(document, chunks)
            if self._document_store is not self._memory_metadata_store:
                self._memory_metadata_store.save_document(document, chunks)
            vector_mutated = False
            job = self._transition_index_job(job=job, status="active")
            stored_document = replace(
                stored_document,
                last_index_status="active",
                last_index_error=None,
            )
            return IngestedDocument(
                knowledge_base_id=document.knowledge_base_id,
                document_id=document.id,
                filename=document.filename,
                chunk_count=len(chunks),
                index_job_id=job.id,
                index_status=job.status,
                document=stored_document,
            )
        except Exception as exc:
            if vector_mutated and document is not None:
                try:
                    if vector_snapshot is not None:
                        self._vector_store.restore_document(
                            document_id=document.id,
                            snapshot=vector_snapshot,
                        )
                    else:
                        self._vector_store.delete_document(document_id=document.id)
                except Exception:
                    pass
            if current_status not in {"active", "failed"}:
                try:
                    self._index_job_store.transition_index_job(
                        job_id=job.id,
                        expected_status=current_status,
                        status="failed",
                        document_id=document.id if document is not None else None,
                        chunk_count=len(chunks) if chunks else None,
                        error=_index_error_message(exc),
                    )
                except Exception:
                    pass
            raise

    def _transition_index_job(
        self,
        *,
        job: IndexJob,
        status: str,
        document_id: str | None = None,
        chunk_count: int | None = None,
    ) -> IndexJob:
        if status not in INDEX_JOB_TRANSITIONS.get(job.status, set()):
            raise RAGProviderError(
                f"invalid index transition: {job.status} -> {status}"
            )
        return self._index_job_store.transition_index_job(
            job_id=job.id,
            expected_status=job.status,
            status=status,
            document_id=document_id,
            chunk_count=chunk_count,
        )

    def ingest_file(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
        content: bytes,
        source_uri: str | None = None,
        title: str | None = None,
        description: str = "",
        tags: list[str] | None = None,
        media_type: str | None = None,
    ) -> IngestedDocument:
        return self._run_index_job(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
            parse_document=lambda: self._parser.parse_bytes(
                knowledge_base_id=knowledge_base_id,
                filename=filename,
                content=content,
                source_uri=source_uri,
                title=title,
                description=description,
                tags=tags,
                media_type=media_type,
                byte_size=len(content),
            ),
        )

    def replace_file(
        self,
        *,
        document: KnowledgeDocument,
        filename: str,
        content: bytes,
        media_type: str | None,
    ) -> IngestedDocument:
        return self._run_index_job(
            knowledge_base_id=document.knowledge_base_id,
            filename=filename,
            parse_document=lambda: self._parser.parse_bytes(
                knowledge_base_id=document.knowledge_base_id,
                filename=filename,
                content=content,
                source_uri=document.source_uri,
                document_id=document.id,
                title=document.title,
                description=document.description,
                tags=document.tags,
                media_type=media_type,
                byte_size=len(content),
            ),
            replacing_document_id=document.id,
        )

    def find_document_by_filename(
        self,
        *,
        knowledge_base_id: str,
        filename: str,
    ) -> KnowledgeDocument | None:
        return self._document_store.find_document_by_filename(
            knowledge_base_id=knowledge_base_id,
            filename=filename,
        )

    def get_document(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
    ) -> KnowledgeDocument | None:
        return self._document_store.get_document(
            knowledge_base_id=knowledge_base_id,
            document_id=document_id,
        )

    def list_documents(
        self,
        *,
        knowledge_base_id: str,
        query: str,
        status: str | None,
        sort: str,
        page: int,
        page_size: int,
    ) -> tuple[list[KnowledgeDocument], int]:
        return self._document_store.list_documents(
            knowledge_base_id=knowledge_base_id,
            query=query,
            status=status,
            sort=sort,
            page=page,
            page_size=page_size,
        )

    def update_document_metadata(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
        title: str,
        description: str,
        tags: list[str],
    ) -> KnowledgeDocument | None:
        updated = self._document_store.update_document_metadata(
            knowledge_base_id=knowledge_base_id,
            document_id=document_id,
            title=title,
            description=description,
            tags=tags,
        )
        if updated is not None and self._document_store is not self._memory_metadata_store:
            memory_document = self._memory_metadata_store.get_document(
                knowledge_base_id=knowledge_base_id,
                document_id=document_id,
            )
            if memory_document is not None:
                self._memory_metadata_store.update_document_metadata(
                    knowledge_base_id=knowledge_base_id,
                    document_id=document_id,
                    title=title,
                    description=description,
                    tags=tags,
                )
        return updated

    def delete_document(
        self,
        *,
        knowledge_base_id: str,
        document_id: str,
    ) -> bool:
        snapshot = self._vector_store.snapshot_document(document_id=document_id)
        self._vector_store.delete_document(document_id=document_id)
        try:
            deleted = self._document_store.delete_document(
                knowledge_base_id=knowledge_base_id,
                document_id=document_id,
            )
            if not deleted:
                self._vector_store.restore_document(
                    document_id=document_id,
                    snapshot=snapshot,
                )
                return False
            if self._document_store is not self._memory_metadata_store:
                self._memory_metadata_store.delete_document(
                    knowledge_base_id=knowledge_base_id,
                    document_id=document_id,
                )
            return True
        except Exception:
            self._vector_store.restore_document(
                document_id=document_id,
                snapshot=snapshot,
            )
            raise

    def delete_knowledge_base(self, *, knowledge_base_id: str) -> None:
        self._vector_store.delete_knowledge_base(
            knowledge_base_id=knowledge_base_id
        )
        self._memory_metadata_store.delete_knowledge_base(
            knowledge_base_id=knowledge_base_id
        )

    def get_index_job(self, *, job_id: str) -> IndexJob | None:
        return self._index_job_store.get_index_job(job_id)

    def list_index_jobs(
        self,
        *,
        knowledge_base_id: str,
        limit: int = 50,
    ) -> list[IndexJob]:
        return self._index_job_store.list_index_jobs(
            knowledge_base_id=knowledge_base_id,
            limit=limit,
        )

    def search(
        self,
        *,
        knowledge_base_id: str,
        query: str,
        limit: int = 5,
        recall_limit: int | None = None,
        rerank_enabled: bool | None = None,
        retrieval_mode: str = "hybrid",
    ) -> list[RetrievedDocument]:
        return self.search_with_metadata(
            knowledge_base_id=knowledge_base_id,
            query=query,
            limit=limit,
            recall_limit=recall_limit,
            rerank_enabled=rerank_enabled,
            retrieval_mode=retrieval_mode,
        ).results

    def reranker_capabilities(self) -> RerankerCapabilities:
        return RerankerCapabilities(
            available=self._reranker_available,
            provider=self._reranker_provider,
            model=self._reranker_model,
            default_enabled=self._rerank_default_enabled,
            status=(
                str(getattr(self._reranker, "status", "ready"))
                if self._reranker_available
                else "unavailable"
            ),
        )

    def search_with_metadata(
        self,
        *,
        knowledge_base_id: str,
        query: str,
        limit: int = 5,
        recall_limit: int | None = None,
        rerank_enabled: bool | None = None,
        retrieval_mode: str = "hybrid",
    ) -> RAGSearchResult:
        query = _normalize_text(query)
        if not query:
            raise RAGValidationError("query is empty")
        retrieval_mode = _normalize_retrieval_mode(retrieval_mode)
        rerank_requested = (
            self._rerank_default_enabled
            if rerank_enabled is None
            else rerank_enabled
        )
        if rerank_requested and not self._reranker_available:
            raise RAGRerankerUnavailableError(
                "reranker is not configured on this server"
            )
        candidate_limit = recall_limit or self._default_recall_limit
        candidate_limit = max(candidate_limit, limit)
        dense_candidates: list[RetrievedDocument] = []
        if retrieval_mode != "lexical":
            with model_usage_scope(
                operation="embedding",
                resource_id=knowledge_base_id,
            ):
                query_embedding = self._embedding_provider.embed_texts(
                    [query],
                    task_type="query",
                )[0]
            dense_candidates = self._vector_store.search(
                knowledge_base_id=knowledge_base_id,
                query_embedding=query_embedding,
                limit=candidate_limit,
            )
        lexical_search = getattr(self._document_store, "search_lexical", None)
        lexical_candidates: list[RetrievedDocument] = []
        if retrieval_mode != "dense":
            lexical_candidates = (
                lexical_search(
                    knowledge_base_id=knowledge_base_id,
                    query=query,
                    limit=candidate_limit,
                )
                if callable(lexical_search)
                else self._memory_metadata_store.search_lexical(
                    knowledge_base_id=knowledge_base_id,
                    query=query,
                    limit=candidate_limit,
                )
            )
        candidates = _rank_retrieval_candidates(
            retrieval_mode=retrieval_mode,
            dense_candidates=dense_candidates,
            lexical_candidates=lexical_candidates,
            lexical_weight=self._lexical_weight,
            rrf_k=self._rrf_k,
        )
        rerank_duration_ms: float | None = None
        rerank_applied = False
        if rerank_requested:
            started_at = perf_counter()
            results = self._reranker.rerank(
                query=query,
                candidates=candidates,
                limit=limit,
            )
            rerank_duration_ms = round((perf_counter() - started_at) * 1000, 3)
            rerank_applied = bool(candidates)
        else:
            results = candidates[:limit]
        return RAGSearchResult(
            results=results,
            retrieval=RetrievalExecution(
                rerank_requested=rerank_requested,
                rerank_applied=rerank_applied,
                provider=self._reranker_provider if rerank_requested else None,
                model=self._reranker_model if rerank_requested else None,
                candidate_count=len(candidates),
                result_count=len(results),
                rerank_duration_ms=rerank_duration_ms,
                retrieval_mode=retrieval_mode,
            ),
        )

    def answer_question(
        self,
        *,
        knowledge_base_id: str,
        question: str,
        llm_client: LLMClient,
        provider: str | None = None,
        model: str | None = None,
        thinking_level: str | None = None,
        limit: int = 5,
        recall_limit: int | None = None,
        rerank_enabled: bool | None = None,
        retrieval_mode: str = "hybrid",
    ) -> RAGAnswer:
        search_result = self.search_with_metadata(
            knowledge_base_id=knowledge_base_id,
            query=question,
            limit=limit,
            recall_limit=recall_limit,
            rerank_enabled=rerank_enabled,
            retrieval_mode=retrieval_mode,
        )
        citations = search_result.results
        messages = self.build_prompt_messages(question=question, citations=citations)

        answer_parts: list[str] = []
        for event in llm_client.stream_chat(
            messages,
            provider=provider,
            model=model,
            thinking_level=thinking_level,
        ):
            if event.type == "delta":
                answer_parts.append(event.text)
            elif event.type == "done":
                break

        return RAGAnswer(
            answer="".join(answer_parts).strip(),
            citations=citations,
            retrieval=search_result.retrieval,
        )

    def build_prompt_messages(
        self,
        *,
        question: str,
        citations: list[RetrievedDocument],
    ) -> list[dict[str, str]]:
        return build_rag_prompt_messages(
            question=question,
            citations=citations,
            max_prompt_chars=self._max_prompt_chars,
        )

    def _format_context(self, citations: list[RetrievedDocument]) -> str:
        return _format_rag_context(
            citations,
            max_prompt_chars=self._max_prompt_chars,
        )


def build_rag_prompt_messages(
    *,
    question: str,
    citations: list[RetrievedDocument],
    max_prompt_chars: int = 6000,
) -> list[dict[str, str]]:
    """Build the production RAG answer prompt from explicit evidence.

    The pure boundary lets generation evaluations use oracle or adversarial
    evidence without invoking retrieval a second time.
    """

    context = _format_rag_context(citations, max_prompt_chars=max_prompt_chars)
    return [
        {
            "role": "system",
            "content": (
                "你是一个企业知识库问答助手。只能基于参考资料回答；"
                "如果参考资料不足，就明确说不知道。回答时保留引用编号。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"用户问题：\n{question}\n\n"
                f"参考资料：\n{context or '没有检索到相关资料。'}\n\n"
                "请给出简洁答案，并在相关句子后标注引用编号，例如 [1]。"
            ),
        },
    ]


def _format_rag_context(
    citations: list[RetrievedDocument],
    *,
    max_prompt_chars: int,
) -> str:
    if max_prompt_chars <= 0:
        raise ValueError("max_prompt_chars must be positive")
    parts: list[str] = []
    used_chars = 0
    for index, item in enumerate(citations, start=1):
        header = (
            f"[{index}] file={item.filename}; "
            f"document_id={item.document_id}; chunk={item.chunk_index}; "
            f"{_line_range_label(item)}"
            f"{_symbols_label(item)}"
            f"score={item.score:.3f}\n"
        )
        body = item.text.strip()
        block = f"{header}{body}"
        remaining = max_prompt_chars - used_chars
        if remaining <= 0:
            break
        if len(block) > remaining:
            block = block[:remaining].rstrip()
        parts.append(block)
        used_chars += len(block)
    return "\n\n".join(parts)


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf_text(content: bytes) -> str:
    if not content:
        raise RAGValidationError("uploaded file is empty")
    try:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise RAGValidationError("encrypted PDF files are not supported")
        pages = [page.extract_text() or "" for page in reader.pages]
    except RAGValidationError:
        raise
    except Exception as exc:
        raise RAGValidationError("unable to read PDF document") from exc
    return "\n\n".join(pages)


def _extract_docx_text(content: bytes) -> str:
    if not content:
        raise RAGValidationError("uploaded file is empty")
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise RAGValidationError("unable to read DOCX document") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def _normalize_code_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = "\n".join(lines)
    normalized = re.sub(r"\n{4,}", "\n\n\n", normalized)
    return normalized.strip()


def _file_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[1].lower()


def _line_number_for_offset(text: str, offset: int) -> int:
    return text.count("\n", 0, max(offset, 0)) + 1


def _code_symbol_starts(
    lines: list[str],
    filename: str,
) -> list[tuple[int, list[str]]]:
    if _file_extension(filename) == ".py":
        python_starts = _python_symbol_starts(lines)
        if python_starts:
            return python_starts

    starts: list[tuple[int, list[str]]] = []
    for index, line in enumerate(lines, start=1):
        symbol = _code_symbol_name(line, filename)
        if symbol is not None:
            starts.append((index, [symbol]))
    return starts


def _python_symbol_starts(lines: list[str]) -> list[tuple[int, list[str]]]:
    """Return top-level Python blocks with class-qualified child symbols.

    AST boundaries keep a class and its methods in one semantic block. Invalid
    or incomplete Python falls back to the language-neutral regex scanner.
    """

    try:
        tree = ast.parse("\n".join(lines))
    except SyntaxError:
        return []

    starts: list[tuple[int, list[str]]] = []
    symbol_nodes = (ast.ClassDef, ast.FunctionDef, ast.AsyncFunctionDef)
    function_nodes = (ast.FunctionDef, ast.AsyncFunctionDef)
    for node in tree.body:
        if not isinstance(node, symbol_nodes):
            continue
        decorator_lines = [item.lineno for item in node.decorator_list]
        start_line = min([node.lineno, *decorator_lines])
        symbols = [node.name]
        if isinstance(node, ast.ClassDef):
            symbols.extend(
                f"{node.name}.{child.name}"
                for child in node.body
                if isinstance(child, function_nodes)
            )
        starts.append((start_line, symbols))
    return starts


def _code_symbol_name(line: str, filename: str) -> str | None:
    extension = _file_extension(filename)
    patterns = _code_symbol_patterns(extension)
    for pattern in patterns:
        match = re.match(pattern, line)
        if match:
            return str(match.group("name"))
    return None


def _code_symbol_patterns(extension: str) -> list[str]:
    if extension == ".py":
        return [
            r"^\s*(?:async\s+def|def)\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\s*\(",
            r"^\s*class\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\b",
        ]
    if extension in {".js", ".jsx", ".ts", ".tsx"}:
        return [
            r"^\s*(?:export\s+)?(?:async\s+)?function\s+(?P<name>[A-Za-z_$][A-Za-z0-9_$]*)\s*\(",
            r"^\s*(?:export\s+)?class\s+(?P<name>[A-Za-z_$][A-Za-z0-9_$]*)\b",
            r"^\s*(?:export\s+)?(?:const|let|var)\s+(?P<name>[A-Za-z_$][A-Za-z0-9_$]*)\s*=\s*(?:async\s*)?(?:\([^)]*\)|[A-Za-z_$][A-Za-z0-9_$]*)\s*=>",
        ]
    if extension == ".go":
        return [r"^\s*func\s+(?:\([^)]*\)\s*)?(?P<name>[A-Za-z_][A-Za-z0-9_]*)\s*\("]
    if extension == ".rs":
        return [
            r"^\s*(?:pub\s+)?fn\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\s*\(",
            r"^\s*(?:pub\s+)?(?:struct|enum|trait|impl)\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\b",
        ]
    if extension == ".java":
        return [
            r"^\s*(?:public|private|protected)?\s*(?:class|interface|enum)\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\b",
            r"^\s*(?:public|private|protected)?\s*(?:static\s+)?[\w<>\[\]]+\s+(?P<name>[A-Za-z_][A-Za-z0-9_]*)\s*\(",
        ]
    return []


def _line_overlap(lines: list[str], max_chars: int) -> list[str]:
    if max_chars <= 0:
        return []
    overlap: list[str] = []
    used_chars = 0
    for line in reversed(lines):
        line_chars = len(line) + 1
        if overlap and used_chars + line_chars > max_chars:
            break
        overlap.append(line)
        used_chars += line_chars
    overlap.reverse()
    return overlap


def _chroma_chunk_metadata(chunk: DocumentChunk) -> dict[str, str | int]:
    metadata: dict[str, str | int] = {
        "knowledge_base_id": chunk.knowledge_base_id,
        "document_id": chunk.document_id,
        "filename": chunk.filename,
        "chunk_index": chunk.chunk_index,
    }
    if chunk.start_line is not None:
        metadata["start_line"] = chunk.start_line
    if chunk.end_line is not None:
        metadata["end_line"] = chunk.end_line
    if chunk.symbols:
        metadata["symbols"] = ",".join(chunk.symbols)
    return metadata


def _optional_int(value: object) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _optional_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _metadata_symbols(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item) for item in value if str(item)]
    if isinstance(value, str):
        return [item for item in value.split(",") if item]
    return []


def _line_range_label(item: RetrievedDocument) -> str:
    if item.start_line is None or item.end_line is None:
        return ""
    return f"lines={item.start_line}-{item.end_line}; "


def _symbols_label(item: RetrievedDocument) -> str:
    if not item.symbols:
        return ""
    return "symbols=" + ",".join(item.symbols[:3]) + "; "


def _document_id(
    *,
    knowledge_base_id: str,
    filename: str,
    source_uri: str | None,
) -> str:
    key = f"{knowledge_base_id}:{filename}:{source_uri or ''}"
    return f"doc_{uuid5(NAMESPACE_URL, key).hex[:16]}"


def _chunk_id(*, document_id: str, chunk_index: int) -> str:
    return f"chk_{uuid5(NAMESPACE_URL, f'{document_id}:{chunk_index}').hex[:16]}"


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[a-zA-Z0-9_]+|[\u4e00-\u9fff]", text.lower())


def _reciprocal_rank_fusion(
    *,
    dense_candidates: list[RetrievedDocument],
    lexical_candidates: list[RetrievedDocument],
    lexical_weight: float,
    rrf_k: int,
) -> list[RetrievedDocument]:
    """Fuse independent dense and lexical rankings with weighted RRF."""

    dense_weight = 1.0 - lexical_weight
    rows: dict[str, dict[str, object]] = {}
    for rank, candidate in enumerate(dense_candidates, start=1):
        rows[candidate.id] = {
            "candidate": candidate,
            "dense_rank": rank,
            "dense_score": (
                candidate.recall_score
                if candidate.recall_score is not None
                else candidate.score
            ),
        }
    for rank, candidate in enumerate(lexical_candidates, start=1):
        row = rows.setdefault(candidate.id, {"candidate": candidate})
        row["lexical_rank"] = rank
        row["lexical_score"] = (
            candidate.lexical_score
            if candidate.lexical_score is not None
            else candidate.score
        )

    fused: list[RetrievedDocument] = []
    for row in rows.values():
        candidate = row["candidate"]
        assert isinstance(candidate, RetrievedDocument)
        dense_rank = _optional_int(row.get("dense_rank"))
        lexical_rank = _optional_int(row.get("lexical_rank"))
        fusion_score = 0.0
        if dense_rank is not None and dense_weight > 0:
            fusion_score += dense_weight / (rrf_k + dense_rank)
        if lexical_rank is not None and lexical_weight > 0:
            fusion_score += lexical_weight / (rrf_k + lexical_rank)
        fused.append(
            replace(
                candidate,
                score=fusion_score,
                recall_score=_optional_float(row.get("dense_score")),
                lexical_score=_optional_float(row.get("lexical_score")),
                hybrid_score=fusion_score,
                dense_rank=dense_rank,
                lexical_rank=lexical_rank,
                fusion_score=fusion_score,
            )
        )
    fused.sort(
        key=lambda item: (
            item.fusion_score or 0.0,
            -(item.dense_rank or 10**9),
            -(item.lexical_rank or 10**9),
        ),
        reverse=True,
    )
    return fused


def _normalize_retrieval_mode(retrieval_mode: str) -> str:
    normalized = retrieval_mode.strip().lower()
    if normalized not in {"dense", "lexical", "hybrid"}:
        raise RAGValidationError(
            "retrieval_mode must be one of: dense, lexical, hybrid"
        )
    return normalized


def _rank_retrieval_candidates(
    *,
    retrieval_mode: str,
    dense_candidates: list[RetrievedDocument],
    lexical_candidates: list[RetrievedDocument],
    lexical_weight: float,
    rrf_k: int,
) -> list[RetrievedDocument]:
    if retrieval_mode == "dense":
        return [
            replace(
                candidate,
                dense_rank=rank,
                recall_score=(
                    candidate.recall_score
                    if candidate.recall_score is not None
                    else candidate.score
                ),
            )
            for rank, candidate in enumerate(dense_candidates, start=1)
        ]
    if retrieval_mode == "lexical":
        return [
            replace(
                candidate,
                lexical_rank=rank,
                lexical_score=(
                    candidate.lexical_score
                    if candidate.lexical_score is not None
                    else candidate.score
                ),
            )
            for rank, candidate in enumerate(lexical_candidates, start=1)
        ]
    return _reciprocal_rank_fusion(
        dense_candidates=dense_candidates,
        lexical_candidates=lexical_candidates,
        lexical_weight=lexical_weight,
        rrf_k=rrf_k,
    )


def _bm25_rank_chunks(
    *,
    query: str,
    chunks: list[DocumentChunk],
    limit: int,
) -> list[RetrievedDocument]:
    query_tokens = _lexical_tokens(query)
    if not query_tokens or not chunks:
        return []
    document_tokens = [
        _lexical_tokens(
            " ".join(
                [
                    chunk.filename,
                    " ".join(chunk.symbols),
                    " ".join(chunk.symbols),
                    chunk.text,
                ]
            )
        )
        for chunk in chunks
    ]
    document_frequency = Counter(
        token
        for tokens in document_tokens
        for token in set(tokens)
    )
    average_length = (
        sum(len(tokens) for tokens in document_tokens) / len(document_tokens)
    )
    scored: list[RetrievedDocument] = []
    for chunk, tokens in zip(chunks, document_tokens):
        if not tokens:
            continue
        frequencies = Counter(tokens)
        score = 0.0
        for term in query_tokens:
            frequency = frequencies.get(term, 0)
            if not frequency:
                continue
            doc_frequency = document_frequency[term]
            inverse_document_frequency = math.log(
                1.0
                + (
                    len(chunks) - doc_frequency + 0.5
                )
                / (doc_frequency + 0.5)
            )
            length_normalization = frequency + 1.5 * (
                1.0 - 0.75
                + 0.75 * len(tokens) / max(average_length, 1.0)
            )
            score += inverse_document_frequency * (
                frequency * 2.5 / length_normalization
            )
        if score <= 0:
            continue
        scored.append(
            RetrievedDocument(
                id=chunk.id,
                knowledge_base_id=chunk.knowledge_base_id,
                document_id=chunk.document_id,
                filename=chunk.filename,
                chunk_index=chunk.chunk_index,
                text=chunk.text,
                score=score,
                start_line=chunk.start_line,
                end_line=chunk.end_line,
                symbols=chunk.symbols,
                lexical_score=score,
            )
        )
    scored.sort(key=lambda item: (item.score, item.id), reverse=True)
    return scored[:limit]


def _lexical_tokens(text: str) -> list[str]:
    tokens: list[str] = []
    for raw_term in _tokenize(text):
        tokens.append(raw_term)
        tokens.extend(
            part
            for part in raw_term.split("_")
            if part and part != raw_term
        )
        tokens.extend(
            part.lower()
            for part in re.findall(
                r"[A-Z]?[a-z]+|[A-Z]+(?=[A-Z]|$)|\d+",
                raw_term,
            )
            if part.lower() != raw_term
        )
    return tokens


def _embedding_text_tokens(text: str) -> int:
    return len(_tokenize(text))


def _embedding_usage_context() -> UsageContext:
    current = current_model_usage_context()
    return UsageContext(
        session_id=current.session_id,
        workspace_id=current.workspace_id,
        operation="embedding",
        resource_id=current.resource_id,
    )


def _gemini_task_type(task_type: str) -> str:
    if task_type == "query":
        return "RETRIEVAL_QUERY"
    return "RETRIEVAL_DOCUMENT"


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right:
        return 0.0
    return sum(a * b for a, b in zip(left, right))


def _first_result_list(value: object) -> list:
    if isinstance(value, list) and value:
        first = value[0]
        if isinstance(first, list):
            return first
    return []


def _qdrant_model(name: str):
    try:
        from qdrant_client import models
    except ImportError as exc:
        raise RAGConfigurationError(
            "qdrant-client is not installed; run pip install -r requirements.txt"
        ) from exc
    return getattr(models, name)


def _qdrant_point_id(chunk_id: str) -> str:
    return str(uuid5(NAMESPACE_URL, chunk_id))


def _qdrant_vector_size(collection_info: object) -> int | None:
    config = getattr(collection_info, "config", None)
    params = getattr(config, "params", None)
    vectors = getattr(params, "vectors", None)
    if vectors is None:
        return None
    size = getattr(vectors, "size", None)
    if size is not None:
        return int(size)
    if isinstance(vectors, dict) and vectors:
        first_vector = next(iter(vectors.values()))
        named_size = getattr(first_vector, "size", None)
        return int(named_size) if named_size is not None else None
    return None


def _validate_embeddings(
    chunks: list[DocumentChunk],
    embeddings: list[list[float]],
) -> None:
    if len(chunks) != len(embeddings):
        raise RAGProviderError(
            "embedding provider returned a different number of vectors than chunks"
        )
    dimensions = {len(embedding) for embedding in embeddings}
    if chunks and (dimensions == {0} or len(dimensions) != 1):
        raise RAGProviderError("embedding provider returned invalid vector dimensions")


def _supports_index_job_store(store: object) -> bool:
    return store is not None and all(
        callable(getattr(store, name, None))
        for name in (
            "create_index_job",
            "transition_index_job",
            "get_index_job",
            "list_index_jobs",
        )
    )


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _index_error_message(exc: Exception) -> str:
    message = str(exc).strip() or exc.__class__.__name__
    return message[:2000]

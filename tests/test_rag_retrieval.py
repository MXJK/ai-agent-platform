from dataclasses import replace
from io import BytesIO
import sys
from types import SimpleNamespace
import unittest
from unittest.mock import patch

from docx import Document

from ai_agent_platform.integrations.rag import (
    InMemoryVectorStore,
    NoopReranker,
    ParsedDocument,
    RAGService,
    RAGRerankerUnavailableError,
    RAGValidationError,
    RecursiveCharacterChunker,
    SentenceTransformerEmbeddingProvider,
    SentenceTransformerCrossEncoderReranker,
    TextDocumentParser,
    RetrievedDocument,
    evaluate_retrieval,
)
from ai_agent_platform.integrations.rag.service import InMemoryRAGMetadataStore


def _simple_pdf(text: str) -> bytes:
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(body)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(payload)


class ConstantEmbeddingProvider:
    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        return [[1.0, 0.0] for _ in texts]


class MutableEmbeddingProvider(ConstantEmbeddingProvider):
    fail = False

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        if self.fail and task_type == "document":
            raise RuntimeError("embedding outage")
        return super().embed_texts(texts, task_type=task_type)


class CountingEmbeddingProvider(ConstantEmbeddingProvider):
    def __init__(self) -> None:
        self.query_calls = 0

    def embed_texts(
        self,
        texts: list[str],
        *,
        task_type: str = "document",
    ) -> list[list[float]]:
        if task_type == "query":
            self.query_calls += 1
        return super().embed_texts(texts, task_type=task_type)


class CountingVectorStore(InMemoryVectorStore):
    def __init__(self) -> None:
        super().__init__()
        self.search_calls = 0

    def search(self, **kwargs) -> list[RetrievedDocument]:
        self.search_calls += 1
        return super().search(**kwargs)


class CountingDocumentStore(InMemoryRAGMetadataStore):
    def __init__(self) -> None:
        super().__init__()
        self.lexical_search_calls = 0

    def search_lexical(self, **kwargs) -> list[RetrievedDocument]:
        self.lexical_search_calls += 1
        return super().search_lexical(**kwargs)


class DenseDistractorVectorStore(InMemoryVectorStore):
    def search(
        self,
        *,
        knowledge_base_id: str,
        query_embedding: list[float],
        limit: int,
    ) -> list[RetrievedDocument]:
        results = super().search(
            knowledge_base_id=knowledge_base_id,
            query_embedding=query_embedding,
            limit=100,
        )
        return [item for item in results if item.filename == "distractor.md"][:limit]


class RecordingIndexJobStore(InMemoryRAGMetadataStore):
    def __init__(self) -> None:
        super().__init__()
        self.transitions: list[tuple[str, str]] = []

    def transition_index_job(self, **kwargs):
        self.transitions.append((kwargs["expected_status"], kwargs["status"]))
        return super().transition_index_job(**kwargs)


class RecordingReranker:
    status = "ready"

    def __init__(self) -> None:
        self.calls = 0

    def rerank(
        self,
        *,
        query: str,
        candidates: list[RetrievedDocument],
        limit: int,
    ) -> list[RetrievedDocument]:
        self.calls += 1
        return [
            replace(candidate, score=0.9, rerank_score=0.9)
            for candidate in reversed(candidates)
        ][:limit]


class RAGRetrievalTests(unittest.TestCase):
    def test_parser_extracts_pdf_docx_and_utf8_markdown(self) -> None:
        parser = TextDocumentParser()
        markdown = parser.parse_bytes(
            knowledge_base_id="docs",
            filename="guide.md",
            content=b"\xef\xbb\xbf# Falcon\nOffline mode",
        )

        word_buffer = BytesIO()
        word_document = Document()
        word_document.add_heading("Falcon handbook", level=1)
        word_document.add_paragraph("DOCX setup instructions")
        word_document.save(word_buffer)
        word = parser.parse_bytes(
            knowledge_base_id="docs",
            filename="handbook.docx",
            content=word_buffer.getvalue(),
        )
        pdf = parser.parse_bytes(
            knowledge_base_id="docs",
            filename="manual.pdf",
            content=_simple_pdf("Falcon PDF guide"),
        )

        self.assertIn("Offline mode", markdown.text)
        self.assertIn("DOCX setup instructions", word.text)
        self.assertIn("Falcon PDF guide", pdf.text)

    def test_python_ast_keeps_class_methods_in_a_qualified_symbol_block(self) -> None:
        document = ParsedDocument(
            id="doc_1",
            knowledge_base_id="repo",
            filename="services.py",
            text=(
                "@registered\n"
                "class PaymentService:\n"
                "    async def issue_refund(self):\n"
                "        return 'ok'\n\n"
                "def healthcheck():\n"
                "    return True"
            ),
        )

        chunks = RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50).split(
            document
        )

        class_chunk = next(
            chunk for chunk in chunks if "class PaymentService" in chunk.text
        )
        self.assertEqual(class_chunk.start_line, 1)
        self.assertIn("PaymentService", class_chunk.symbols)
        self.assertIn("PaymentService.issue_refund", class_chunk.symbols)
        self.assertFalse(any(chunk.symbols == ["issue_refund"] for chunk in chunks))
        function_chunk = next(chunk for chunk in chunks if "def healthcheck" in chunk.text)
        self.assertEqual(function_chunk.symbols, ["healthcheck"])

    def test_hybrid_retrieval_promotes_exact_symbol_match(self) -> None:
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
            lexical_weight=0.35,
        )
        service.ingest_document(
            knowledge_base_id="repo",
            filename="payments.py",
            content=(
                "class PaymentService:\n"
                "    def issue_refund(self):\n"
                "        return 'refunded'"
            ),
        )
        service.ingest_document(
            knowledge_base_id="repo",
            filename="shipping.py",
            content=(
                "class ShippingService:\n"
                "    def dispatch_order(self):\n"
                "        return 'sent'"
            ),
        )

        results = service.search(
            knowledge_base_id="repo",
            query="PaymentService issue_refund",
            limit=2,
        )

        self.assertEqual(results[0].filename, "payments.py")
        self.assertGreater(results[0].lexical_score or 0.0, 0.0)
        self.assertIsNotNone(results[0].hybrid_score)
        self.assertEqual(results[0].score, results[0].hybrid_score)
        self.assertEqual(results[0].dense_rank, 1)
        self.assertEqual(results[0].lexical_rank, 1)
        self.assertEqual(results[0].fusion_score, results[0].score)

    def test_request_can_enable_or_skip_configured_reranker(self) -> None:
        reranker = RecordingReranker()
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
            reranker=reranker,
            reranker_provider="sentence_transformer",
            reranker_model="test-cross-encoder",
            rerank_default_enabled=False,
            default_recall_limit=10,
            max_prompt_chars=2000,
        )
        service.ingest_document(
            knowledge_base_id="docs",
            filename="guide.md",
            content="request-scoped reranking reference",
        )

        skipped = service.search_with_metadata(
            knowledge_base_id="docs",
            query="reranking",
            limit=1,
            rerank_enabled=False,
        )
        applied = service.search_with_metadata(
            knowledge_base_id="docs",
            query="reranking",
            limit=1,
            rerank_enabled=True,
        )

        self.assertEqual(reranker.calls, 1)
        self.assertFalse(skipped.retrieval.rerank_applied)
        self.assertIsNone(skipped.results[0].rerank_score)
        self.assertTrue(applied.retrieval.rerank_requested)
        self.assertTrue(applied.retrieval.rerank_applied)
        self.assertEqual(applied.retrieval.provider, "sentence_transformer")
        self.assertEqual(applied.retrieval.model, "test-cross-encoder")
        self.assertEqual(applied.results[0].rerank_score, 0.9)
        self.assertIsNotNone(applied.retrieval.rerank_duration_ms)

    def test_request_rejects_unavailable_reranker(self) -> None:
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
            rerank_default_enabled=False,
        )

        with self.assertRaises(RAGRerankerUnavailableError):
            service.search_with_metadata(
                knowledge_base_id="docs",
                query="reranking",
                limit=1,
                rerank_enabled=True,
            )

    def test_cross_encoder_loads_lazily_and_reuses_model(self) -> None:
        created_models = []

        class FakeCrossEncoder:
            def __init__(self, model_name: str, *, device: str) -> None:
                self.model_name = model_name
                self.device = device
                self.predict_calls = 0
                created_models.append(self)

            def predict(self, pairs):
                self.predict_calls += 1
                return [float(index) for index, _ in enumerate(pairs, start=1)]

        reranker = SentenceTransformerCrossEncoderReranker(
            model_name="test-cross-encoder",
            device=" cpu ",
        )
        candidates = [
            RetrievedDocument(
                id="chunk_1",
                knowledge_base_id="docs",
                document_id="doc_1",
                filename="guide.md",
                chunk_index=0,
                text="first",
                score=0.2,
            ),
            RetrievedDocument(
                id="chunk_2",
                knowledge_base_id="docs",
                document_id="doc_1",
                filename="guide.md",
                chunk_index=1,
                text="second",
                score=0.1,
            ),
        ]
        self.assertEqual(reranker.status, "not_loaded")

        with patch.dict(
            sys.modules,
            {
                "sentence_transformers": SimpleNamespace(
                    CrossEncoder=FakeCrossEncoder
                )
            },
        ):
            first = reranker.rerank(query="query", candidates=candidates, limit=2)
            second = reranker.rerank(query="query", candidates=candidates, limit=1)

        self.assertEqual(reranker.status, "ready")
        self.assertEqual(len(created_models), 1)
        self.assertEqual(created_models[0].model_name, "test-cross-encoder")
        self.assertEqual(created_models[0].device, "cpu")
        self.assertEqual(created_models[0].predict_calls, 2)
        self.assertEqual([item.id for item in first], ["chunk_2", "chunk_1"])
        self.assertEqual(second[0].rerank_score, 2.0)

    def test_sentence_transformer_embedding_loads_lazily_and_normalizes(self) -> None:
        created_models = []

        class FakeSentenceTransformer:
            def __init__(self, model_name: str, *, device: str) -> None:
                self.model_name = model_name
                self.device = device
                self.encode_calls = []
                created_models.append(self)

            def encode(self, texts, **kwargs):
                self.encode_calls.append((list(texts), kwargs))
                return [[0.6, 0.8] for _ in texts]

        provider = SentenceTransformerEmbeddingProvider(
            model_name="BAAI/bge-m3",
            device=" cpu ",
        )
        self.assertEqual(provider.status, "not_loaded")

        with patch.dict(
            sys.modules,
            {
                "sentence_transformers": SimpleNamespace(
                    SentenceTransformer=FakeSentenceTransformer
                )
            },
        ):
            first = provider.embed_texts(["中文问题", "English document"])
            second = provider.embed_texts(["reuse"])

        self.assertEqual(provider.status, "ready")
        self.assertEqual(len(created_models), 1)
        self.assertEqual(created_models[0].model_name, "BAAI/bge-m3")
        self.assertEqual(created_models[0].device, "cpu")
        self.assertEqual(len(created_models[0].encode_calls), 2)
        self.assertTrue(
            created_models[0].encode_calls[0][1]["normalize_embeddings"]
        )
        self.assertEqual(first, [[0.6, 0.8], [0.6, 0.8]])
        self.assertEqual(second, [[0.6, 0.8]])

    def test_retrieval_modes_skip_unused_recall_paths(self) -> None:
        embedding_provider = CountingEmbeddingProvider()
        vector_store = CountingVectorStore()
        document_store = CountingDocumentStore()
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=embedding_provider,
            vector_store=vector_store,
            document_store=document_store,
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
        )
        service.ingest_document(
            knowledge_base_id="docs",
            filename="guide.md",
            content="QUASAR multilingual retrieval guide",
        )

        lexical = service.search_with_metadata(
            knowledge_base_id="docs",
            query="QUASAR",
            retrieval_mode="lexical",
        )
        self.assertEqual(embedding_provider.query_calls, 0)
        self.assertEqual(vector_store.search_calls, 0)
        self.assertEqual(document_store.lexical_search_calls, 1)
        self.assertEqual(lexical.retrieval.retrieval_mode, "lexical")
        self.assertIsNotNone(lexical.results[0].lexical_rank)
        self.assertIsNone(lexical.results[0].dense_rank)

        dense = service.search_with_metadata(
            knowledge_base_id="docs",
            query="QUASAR",
            retrieval_mode="dense",
        )
        self.assertEqual(embedding_provider.query_calls, 1)
        self.assertEqual(vector_store.search_calls, 1)
        self.assertEqual(document_store.lexical_search_calls, 1)
        self.assertEqual(dense.retrieval.retrieval_mode, "dense")
        self.assertIsNotNone(dense.results[0].dense_rank)
        self.assertIsNone(dense.results[0].lexical_rank)

        hybrid = service.search_with_metadata(
            knowledge_base_id="docs",
            query="QUASAR",
            retrieval_mode="hybrid",
        )
        self.assertEqual(embedding_provider.query_calls, 2)
        self.assertEqual(vector_store.search_calls, 2)
        self.assertEqual(document_store.lexical_search_calls, 2)
        self.assertEqual(hybrid.retrieval.retrieval_mode, "hybrid")
        self.assertIsNotNone(hybrid.results[0].fusion_score)

        with self.assertRaisesRegex(RAGValidationError, "retrieval_mode"):
            service.search(
                knowledge_base_id="docs",
                query="QUASAR",
                retrieval_mode="unknown",
            )

    def test_lexical_recall_adds_candidate_missing_from_dense_recall(self) -> None:
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=DenseDistractorVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=1,
            max_prompt_chars=2000,
            lexical_weight=0.7,
        )
        service.ingest_document(
            knowledge_base_id="docs",
            filename="distractor.md",
            content="General retention overview without a control identifier.",
        )
        service.ingest_document(
            knowledge_base_id="docs",
            filename="target.md",
            content="The exact control is QUASAR_RETENTION_47X.",
        )

        results = service.search(
            knowledge_base_id="docs",
            query="QUASAR_RETENTION_47X",
            limit=1,
            recall_limit=1,
        )

        self.assertEqual(results[0].filename, "target.md")
        self.assertIsNone(results[0].dense_rank)
        self.assertEqual(results[0].lexical_rank, 1)

    def test_successful_index_job_follows_all_state_transitions(self) -> None:
        job_store = RecordingIndexJobStore()
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
            index_job_store=job_store,
        )

        service.ingest_document(
            knowledge_base_id="docs",
            filename="guide.md",
            content="state machine reference",
        )

        self.assertEqual(
            job_store.transitions,
            [
                ("pending", "parsing"),
                ("parsing", "embedding"),
                ("embedding", "vector_written"),
                ("vector_written", "active"),
            ],
        )

    def test_retrieval_metrics_compute_recall_and_mrr(self) -> None:
        metrics = evaluate_retrieval(
            rankings=[["b.py", "a.py"], ["c.py", "d.py"]],
            relevant_documents=[{"a.py"}, {"missing.py"}],
            k=2,
        )

        self.assertEqual(metrics.evaluated_cases, 2)
        self.assertEqual(metrics.recall_at_k, 0.5)
        self.assertEqual(metrics.mean_reciprocal_rank, 0.25)

    def test_failed_reindex_records_job_and_preserves_previous_vectors(self) -> None:
        embedding_provider = MutableEmbeddingProvider()
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=embedding_provider,
            vector_store=InMemoryVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
        )
        first = service.ingest_document(
            knowledge_base_id="docs",
            filename="guide.md",
            content="stable previous content",
        )
        embedding_provider.fail = True

        with self.assertRaisesRegex(RuntimeError, "embedding outage"):
            service.ingest_document(
                knowledge_base_id="docs",
                filename="guide.md",
                content="replacement content",
            )

        jobs = service.list_index_jobs(knowledge_base_id="docs")
        self.assertEqual([job.status for job in jobs], ["failed", "active"])
        self.assertEqual(jobs[0].error, "embedding outage")
        self.assertEqual(first.index_status, "active")
        embedding_provider.fail = False
        results = service.search(
            knowledge_base_id="docs",
            query="stable previous content",
            limit=1,
        )
        self.assertIn("stable previous content", results[0].text)

    def test_delete_knowledge_base_keeps_other_namespaces(self) -> None:
        service = RAGService(
            parser=TextDocumentParser(),
            chunker=RecursiveCharacterChunker(chunk_size=500, chunk_overlap=50),
            embedding_provider=ConstantEmbeddingProvider(),
            vector_store=InMemoryVectorStore(),
            reranker=NoopReranker(),
            default_recall_limit=10,
            max_prompt_chars=2000,
        )
        for knowledge_base_id in ("delete_me", "keep_me"):
            service.ingest_document(
                knowledge_base_id=knowledge_base_id,
                filename=f"{knowledge_base_id}.md",
                content=f"{knowledge_base_id} reference content",
            )

        service.delete_knowledge_base(knowledge_base_id="delete_me")

        self.assertEqual(
            service.search(
                knowledge_base_id="delete_me",
                query="reference",
                limit=5,
            ),
            [],
        )
        self.assertEqual(
            len(
                service.search(
                    knowledge_base_id="keep_me",
                    query="reference",
                    limit=5,
                )
            ),
            1,
        )


if __name__ == "__main__":
    unittest.main()
